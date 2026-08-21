import streamlit as st
from google import genai
from docx import Document
from fpdf import FPDF
import io
from datetime import date

# 1. Configuración de la API (Pega tu clave AIzaSy...)
client = genai.Client(api_key="AQ.Ab8RN6JM_YMzUguhKme19dTI9laFp2pEaKDzojA5eEQR6nqrRw")

st.set_page_config(page_title="Instrumento de Registro BAP", page_icon="📝", layout="wide")

# --- ESTILOS CSS ---
st.markdown("""
<style>
    .caja-datos { background-color: #F6F1E9; border: 2px solid #C4A462; border-radius: 15px; padding: 25px; margin-bottom: 20px; color: #333333; }
    .caja-instrucciones { border: 2px solid #2B5B41; padding: 20px; margin-bottom: 20px; background-color: #ffffff; color: #333333; }
    .titulo-seccion { color: #2B5B41; font-weight: bold; margin-top: 30px; margin-bottom: 15px;}
    .stCheckbox { margin-bottom: -10px; } 
    hr.separador-movil { border: 0; border-top: 1px solid #e0e0e0; margin: 15px 0; }
</style>
""", unsafe_allow_html=True)

# --- FUNCIONES DE EXPORTACIÓN ---
def generar_docx(texto_informe, alumno):
    doc = Document()
    doc.add_heading('Informe de Análisis y Estrategias BAP', 0)
    doc.add_heading(f'Alumno(a): {alumno}', 1)
    doc.add_paragraph(texto_informe)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_pdf(texto_informe, alumno):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    texto_limpio = texto_informe.encode('latin-1', 'replace').decode('latin-1')
    pdf.cell(200, 10, txt=f"Informe de Analisis BAP - {alumno}", ln=True, align='C')
    pdf.ln(10)
    pdf.multi_cell(0, 7, txt=texto_limpio)
    return bytes(pdf.output())

# --- ENCABEZADO ---
st.title("Instrumento de registro de las barreras para el aprendizaje y la participación")

# --- DATOS GENERALES ---
st.markdown('<div class="caja-datos">', unsafe_allow_html=True)
st.markdown("<h4 style='text-align: center; color: #333;'>DATOS GENERALES</h4>", unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
centro_escolar = col1.text_input("Centro Escolar:")
subsistema = col2.selectbox("Subsistema:", ["Estatal", "Federal"])
turno = col3.selectbox("Turno:", ["Matutino", "Vespertino"])

col4, col5 = st.columns(2)
grado = col4.selectbox("Grado:", ["1°", "2°", "3°", "4°", "5°", "6°"])
grupo = col5.selectbox("Grupo:", ["A", "B", "C"])

nivel = st.radio("Nivel:", ["Inicial", "Preescolar", "Primaria", "Secundaria", "Laboral"], horizontal=True)
nombre_alumno = st.text_input("Nombre de la o el estudiante (NNAJ):")
nombre_docentes = st.text_area("Nombre y función del equipo itinerante, docentes y agentes educativos:")
fecha_eval = st.date_input("Fecha:", value=date.today())
st.markdown('</div>', unsafe_allow_html=True)

# --- CLASIFICACIÓN DE BARRERAS ---
st.markdown("<h3 class='titulo-seccion'>1. Clasificación de Barreras</h3>", unsafe_allow_html=True)
tab1, tab2, tab3, tab4 = st.tabs(["🏛️ Estructurales", "📜 Normativas", "📚 Didácticas", "🤝 Actitudinales"])

bap_seleccionadas = []

with tab1:
    st.write("**BARRERAS ESTRUCTURALES: Normalizan la exclusión**")
    est_list = [
        "Negar la atención educativa por falta de recursos humanos, materiales, apoyos, capacitación, entre otras.",
        "Priorizar una condición específica sobre otra por considerarla más importante para la atención.",
        "Separación de NNAJ para su atención exclusiva o permanente en aulas especiales.",
        "Negar la atención educativa por problemas de administración, vinculación y/o gestión del tiempo laboral.",
        "Infraestructura inadecuada (escolar o en el aula), falta de rampas, accesos, adecuaciones arquitectónicas, señalamientos, entre otras.",
        "Invisibilizar e ignorar la presencia de alguna o algún NNAJ por su condición (no identificación oportuna).",
        "No incluir en el diagnóstico escolar la presencia de NNAJ asociado a una condición específica.",
        "Medios de transporte insuficientes, inadecuados, inaccesibles, etc.",
        "Cuando por su condición, alguna o algún NNAJ asiste a un Centro de Atención Múltiple (CAM) y podría incluirse en un aula o una sala de educación básica."
    ]
    for i, b in enumerate(est_list):
        if st.checkbox(b, key=f"est_{i}"): bap_seleccionadas.append(f"{b} (Estructural)")

with tab2:
    st.write("**BARRERAS NORMATIVAS: Impedimento desde la Ley, norma, disposición política, etc.**")
    nor_list = [
        "Desconocimiento de los documentos normativos que regulan la atención educativa de la diversidad en el marco de la inclusión.",
        "Omitir el cumplimiento, ejecución y difusión de los diferentes documentos normativos.",
        "Falta de políticas compensatorias para el apoyo de la población asociada a una condición específica como becas, programas, recursos, entre otros.",
        "Procesos de gestión excluyentes, descontextualizados, segregatorios, entre otros.",
        "Organización incompleta del centro de trabajo escolar que dificulta la atención a la diversidad.",
        "Falta de un proyecto escolar o del Centro, que incluya a toda la comunidad educativa independientemente de su condición.",
        "Falta de un Programa Analítico que contextualice los Contenidos y Procesos de desarrollo de aprendizaje acordes a la diversidad de la población.",
        "Desvinculación entre los Servicios de Educación Especial (SEE), las escuelas o Centros de Educación Básica, los especialistas, los niveles educativos, el campo laboral, etc.",
        "Falta de mecanismos de formación, profesionalización docente, actualización y/o acompañamiento pedagógico.",
        "Rigidez administrativa, burocrática, morosa, entre otras.",
        "Falta de un liderazgo compartido, colegiado, cooperativo.",
        "Canalización de NNAJ a centros paralelos de atención como única alternativa, no como complemento formativo y de desarrollo."
    ]
    for i, b in enumerate(nor_list):
        if st.checkbox(b, key=f"nor_{i}"): bap_seleccionadas.append(f"{b} (Normativa)")

with tab3:
    st.write("**BARRERAS DIDÁCTICAS: Métodos de enseñanza y evaluación que no son acordes a las necesidades reales de NNAJ.**")
    did_list = [
        "Práctica docente homogeneizada sin considerar la diversidad y las necesidades específicas de cada condición.",
        "No diversificar la práctica docente y la planeación didáctica conforme a los Principios del Diseño Universal para el Aprendizaje (DUA).",
        "Mecanismos de evaluación homogéneos sin considerar la diversidad de NNAJ, sus capacidades y/o necesidades.",
        "Ausencia de Ajustes Razonables (AR) acordes a las necesidades específicas asociadas a la condición (cuando se requieran).",
        "Desconocimiento del qué, para qué, cómo y cuándo enseñar y/o evaluar.",
        "Estandarización de los aprendizajes como punto de partida para la intervención y para la evaluación.",
        "Priorizar el trabajo individualizado sobre el trabajo cooperativo y colaborativo dentro del aula.",
        "Desvinculación del trabajo de enseñanza y aprendizaje entre docentes de educación básica, y el personal de los equipos de apoyo de educación especial.",
        "Desvinculación del trabajo de enseñanza y aprendizaje con madres padres de familia y tutores de NNAJ.",
        "Ausencia de recursos y/o apoyos tecnológicos como estrategia de enriquecimiento o una medida compensatoria."
    ]
    for i, b in enumerate(did_list):
        if st.checkbox(b, key=f"did_{i}"): bap_seleccionadas.append(f"{b} (Didáctica)")

with tab4:
    st.write("**BARRERAS ACTITUDINALES: Relacionadas con las interacciones y concepciones sociales.**")
    st.info("📱 Selecciona los contextos en los que se hace evidente cada barrera.")
    
    act_list = [
        "Apatía, rechazo o indiferencia hacia las condiciones específicas.",
        "Segregación y/o exclusión en los procesos de inclusión.",
        "Sobreprotección.",
        "Altas expectativas sobre su desarrollo y aprendizaje.",
        "Bajas expectativas sobre su desarrollo y aprendizaje.",
        "Desconocimiento ante las condiciones específicas que tienen NNAJ.",
        "Acoso y/o bullying en el entorno.",
        "Estereotipos erróneos y/o prejuicios ante las condiciones.",
        "Etiquetas sobre las condiciones que se tienen.",
        "Clasismo y/o racismo generado por la condición personal, social, económica, física o cultural.",
        "Exclusión por rendimiento escolar, desempeño, desarrollo, entre otras.",
        "Abuso físico y/o emocional generado por su condición.",
        "Poco aprecio por la educación o el desarrollo de NNAJ.",
        "Discriminación por género, ideología, creencias, o alguna otra."
    ]
    
    # Diseño Optimizado para Teléfonos Móviles
    for i, barrera in enumerate(act_list):
        st.markdown(f"**{i+1}. {barrera}**")
        c1, c2, c3 = st.columns(3)
        if c1.checkbox("🏫 Aula/Sala", key=f"a_{i}"): bap_seleccionadas.append(f"{barrera} (Actitudinal - Aula)")
        if c2.checkbox("🏢 Escuela/Centro", key=f"e_{i}"): bap_seleccionadas.append(f"{barrera} (Actitudinal - Escuela)")
        if c3.checkbox("👨‍👩‍👧 Familia", key=f"f_{i}"): bap_seleccionadas.append(f"{barrera} (Actitudinal - Familia)")
        st.markdown("<hr class='separador-movil'>", unsafe_allow_html=True)

st.divider()

# --- MÓDULO DE INTERPRETACIÓN CUALITATIVA ---
st.markdown("<h3 class='titulo-seccion'>2. Interpretación y Estrategias (Guía Oficial)</h3>", unsafe_allow_html=True)
st.info("💡 Completa los apartados necesarios para orientar el plan de acción. La Inteligencia Artificial articulará tus respuestas con las barreras seleccionadas para generar el borrador.")

with st.expander("Desplegar formulario de análisis cualitativo", expanded=True):
    col_preg1, col_preg2 = st.columns(2)
    
    with col_preg1:
        q1 = st.text_input("1. ¿En qué contextos están enfrentando BAP?")
        q2 = st.text_input("2. ¿Quiénes son las y/o los agentes educativos que están generando las BAP?")
        q3 = st.text_area("3. ¿Qué tipo de barreras son las que están enfrentando, acorde a la clasificación?")
        q4 = st.text_input("4. Si fuera necesario priorizar ¿con cuáles debería empezar la intervención?")
        q5 = st.text_area("5. ¿Pueden ser minimizadas desde el aula o necesitan un plan complementario?")
        
    with col_preg2:
        q6 = st.text_area("6. ¿Cuáles son las posibles estrategias que se deben realizar?")
        q7 = st.text_input("7. ¿Cuáles serán las responsabilidades de cada integrante del colectivo?")
        q8 = st.text_input("8. ¿En qué plazo se realizarán las acciones?")
        q9 = st.text_area("9. ¿Cómo aportarán estas acciones a los retos en pro de la inclusión?")
    
    q10 = st.text_area("10. Cualquier otro dato que aporte a la interpretación y establecimiento de acciones:")

st.divider()

# --- GENERACIÓN DE INFORME IA ---
if st.button("✨ Generar Informe Potenciado por IA", type="primary"):
    if not bap_seleccionadas and not q1:
        st.warning("Selecciona al menos una barrera o responde alguna pregunta de la guía antes de generar el documento.")
    else:
        with st.spinner("Procesando análisis técnico y estructurando el informe final..."):
            
            prompt = f"""
            Actúa como un especialista en Educación Especial, redactando el "Informe de las BAP que enfrentan niñas, niños, adolescentes y jóvenes".
            
            DATOS DE IDENTIFICACIÓN:
            Alumno: {nombre_alumno} ({nivel} {grado} {grupo}, {centro_escolar}, Subsistema: {subsistema}, Turno: {turno})
            Equipo o docentes a cargo: {nombre_docentes}
            
            1. BARRERAS DETECTADAS (Acorde a la normatividad):
            {', '.join(bap_seleccionadas) if bap_seleccionadas else 'No se marcaron casillas específicas, analizar cualitativamente.'}
            
            2. ANÁLISIS CUALITATIVO Y ESTRATÉGICO DEL DOCENTE:
            - Contextos de las BAP: {q1}
            - Agentes generadores: {q2}
            - Tipos de barreras priorizadas: {q3}
            - Prioridad de intervención: {q4}
            - Alcance (Aula o plan complementario): {q5}
            - Estrategias propuestas: {q6}
            - Responsabilidades del colectivo: {q7}
            - Plazos de acción: {q8}
            - Aportación a la inclusión: {q9}
            - Datos extra relevantes: {q10}
            
            REGLAS DE REDACCIÓN Y ESTRUCTURA:
            No generes un documento de formato "Pregunta/Respuesta". Redacta un informe narrativo formal, profesional y cohesionado, dividido por secciones claras (Introducción, Barreras Identificadas, Estrategias de Intervención, Responsabilidades y Plazos). Integra armónicamente las casillas marcadas con el texto abierto escrito por el docente. Omite cualquier variable o pregunta que el docente haya dejado en blanco, llenando el vacío lógico de forma técnica y pedagógica.
            """
            
            try:
                response = client.models.generate_content(
                    model="gemini-3.6-flash", 
                    contents=prompt
                )
                
                st.success("¡Documento maestro generado con éxito!")
                
                col_w, col_p = st.columns(2)
                with col_w:
                    docx_file = generar_docx(response.text, nombre_alumno)
                    st.download_button("📄 Descargar Archivo Editable (Word)", data=docx_file, file_name=f"Informe_BAP_{nombre_alumno}.docx")
                with col_p:
                    pdf_file = generar_pdf(response.text, nombre_alumno)
                    st.download_button("📕 Descargar PDF Final", data=pdf_file, file_name=f"Informe_BAP_{nombre_alumno}.pdf")
                
                with st.expander("Ver borrador del informe generado", expanded=True):
                    st.write(response.text)
                    
            except Exception as e:
                st.error(f"Error de conexión con la IA: {e}")
